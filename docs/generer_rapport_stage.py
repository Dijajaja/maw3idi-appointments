#!/usr/bin/env python3
"""
Génère un rapport de stage complet en format Word (.docx)
avec toutes les captures d'écran et diagrammes UML intégrés
"""

import os
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Chemins
PROJECT_ROOT = Path(__file__).parent.parent
SCREENSHOTS_DIR = PROJECT_ROOT / "docs" / "screenshots"
UML_DIR = PROJECT_ROOT / "docs" / "uml"
OUTPUT_DIR = PROJECT_ROOT / "docs"
from datetime import datetime
OUTPUT_FILE = OUTPUT_DIR / f"RAPPORT_DE_STAGE_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

def add_page_break(doc):
    """Ajoute un saut de page"""
    doc.add_page_break()

def add_heading_with_style(doc, text, level=1):
    """Ajoute un titre avec style"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_image_to_doc(doc, image_path, caption=None, width_cm=16):
    """Ajoute une image au document avec une légende"""
    if not image_path.exists():
        para = doc.add_paragraph(f"[Image non trouvée: {image_path.name}]")
        para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        return
    
    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_cm / 2.54))
        
        if caption:
            caption_para = doc.add_paragraph(caption, style='Caption')
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.runs[0].font.italic = True
            caption_para.runs[0].font.size = Pt(10)
    except Exception as e:
        para = doc.add_paragraph(f"[Erreur lors de l'ajout de l'image: {str(e)}]")
        para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

def create_custom_styles(doc):
    """Crée des styles personnalisés"""
    styles = doc.styles
    
    # Style pour les légendes
    if 'Caption' not in [s.name for s in styles]:
        caption_style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_font = caption_style.font
        caption_font.size = Pt(10)
        caption_font.italic = True
        caption_font.color.rgb = RGBColor(64, 64, 64)

def generate_rapport():
    """Génère le rapport de stage complet"""
    
    if not DOCX_AVAILABLE:
        print("❌ python-docx n'est pas installé.")
        print("Installez-le avec: pip install python-docx")
        return False
    
    print("=" * 70)
    print("📄 GÉNÉRATION DU RAPPORT DE STAGE")
    print("=" * 70)
    print()
    
    # Créer le document
    doc = Document()
    
    # Définir les marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    create_custom_styles(doc)
    
    # ============================================
    # PAGE DE GARDE
    # ============================================
    print("📝 Création de la page de garde...")
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("RAPPORT DE STAGE")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espace
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Système de Gestion de Rendez-vous")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.name = 'Arial'
    
    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace
    
    tech_para = doc.add_paragraph()
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_run = tech_para.add_run("Application Web Django")
    tech_run.font.size = Pt(14)
    tech_run.font.italic = True
    
    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Année {datetime.now().year}")
    date_run.font.size = Pt(12)
    
    add_page_break(doc)
    
    # ============================================
    # TABLE DES MATIÈRES
    # ============================================
    print("📋 Création de la table des matières...")
    
    doc.add_heading("Table des Matières", 1)
    doc.add_paragraph("1. Introduction", style='List Number')
    doc.add_paragraph("2. Présentation du Projet", style='List Number')
    doc.add_paragraph("3. Architecture Technique", style='List Number')
    doc.add_paragraph("4. Modélisation UML", style='List Number')
    doc.add_paragraph("5. Interface Utilisateur", style='List Number')
    doc.add_paragraph("6. Fonctionnalités Principales", style='List Number')
    doc.add_paragraph("7. Conclusion", style='List Number')
    
    add_page_break(doc)
    
    # ============================================
    # 1. INTRODUCTION
    # ============================================
    print("📝 Section 1: Introduction...")
    
    doc.add_heading("1. Introduction", 1)
    
    intro_text = """
Ce rapport présente le développement d'un système de gestion de rendez-vous basé sur le framework Django. 
L'application permet aux clients de réserver des rendez-vous en ligne pour différents services, avec une 
gestion complète des disponibilités, des paiements et des notifications.

Le système a été conçu pour être flexible, évolutif et facile à utiliser, tant pour les administrateurs 
que pour les clients finaux. Il intègre des fonctionnalités avancées telles que la gestion des horaires 
de travail, les jours de congé, la reprogrammation de rendez-vous, et un système de paiement multi-méthodes.
"""
    
    doc.add_paragraph(intro_text.strip())
    
    add_page_break(doc)
    
    # ============================================
    # 2. PRÉSENTATION DU PROJET
    # ============================================
    print("📝 Section 2: Présentation du projet...")
    
    doc.add_heading("2. Présentation du Projet", 1)
    
    doc.add_heading("2.1. Contexte et Objectifs", 2)
    
    contexte_text = """
Le système de gestion de rendez-vous a été développé pour répondre aux besoins de modernisation de la 
prise de rendez-vous en ligne. L'objectif principal est de permettre aux entreprises et professionnels 
de gérer efficacement leurs rendez-vous tout en offrant une expérience utilisateur optimale.
"""
    doc.add_paragraph(contexte_text.strip())
    
    doc.add_heading("2.2. Technologies Utilisées", 2)
    
    tech_list = [
        "Framework Django 5.2.7 - Framework web Python",
        "Python 3.x - Langage de programmation",
        "SQLite/PostgreSQL - Base de données",
        "HTML/CSS/JavaScript - Interface utilisateur",
        "Playwright - Automatisation des tests et captures",
        "Django Q2 - Gestion des tâches asynchrones",
        "Babel - Internationalisation (i18n)",
        "Pillow - Traitement d'images"
    ]
    
    for tech in tech_list:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading("2.3. Fonctionnalités Principales", 2)
    
    features_list = [
        "Gestion des services avec prix, durée et images",
        "Gestion des membres du personnel avec horaires personnalisés",
        "Système de réservation en ligne avec calendrier interactif",
        "Gestion des disponibilités et conflits de rendez-vous",
        "Système de paiement multi-méthodes (carte, virement, portefeuilles électroniques)",
        "Notifications par email avec rappels automatiques",
        "Reprogrammation de rendez-vous",
        "Interface d'administration complète",
        "Tableau de bord avec statistiques",
        "Support multilingue (français, anglais)"
    ]
    
    for feature in features_list:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_page_break(doc)
    
    # ============================================
    # 3. ARCHITECTURE TECHNIQUE
    # ============================================
    print("📝 Section 3: Architecture technique...")
    
    doc.add_heading("3. Architecture Technique", 1)
    
    doc.add_heading("3.1. Structure du Projet", 2)
    
    structure_text = """
Le projet suit l'architecture MVC (Model-View-Controller) de Django :
- Models : Définition des entités métier (Service, Appointment, StaffMember, etc.)
- Views : Gestion de la logique métier et des requêtes HTTP
- Templates : Interface utilisateur (HTML)
- URLs : Routage des requêtes
- Static : Fichiers statiques (CSS, JavaScript, images)
"""
    doc.add_paragraph(structure_text.strip())
    
    # Diagramme de composants
    if (UML_DIR / "diagramme_composants.png").exists():
        doc.add_heading("3.2. Architecture des Composants", 2)
        add_image_to_doc(
            doc, 
            UML_DIR / "diagramme_composants.png",
            "Figure 1: Architecture des composants du système",
            width_cm=16
        )
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # ============================================
    # 4. MODÉLISATION UML
    # ============================================
    print("📝 Section 4: Diagrammes UML...")
    
    doc.add_heading("4. Modélisation UML", 1)
    
    # Diagramme de classe
    if (UML_DIR / "diagramme_classe.png").exists():
        doc.add_heading("4.1. Diagramme de Classe", 2)
        doc.add_paragraph(
            "Le diagramme de classe présente la structure des modèles Django et leurs relations. "
            "Il montre les 11 modèles principaux du système et leurs interactions."
        )
        add_image_to_doc(
            doc,
            UML_DIR / "diagramme_classe.png",
            "Figure 2: Diagramme de classe - Structure des modèles",
            width_cm=18
        )
        doc.add_paragraph()
        add_page_break(doc)
    
    # Diagramme de cas d'utilisation
    if (UML_DIR / "diagramme_cas_utilisation.png").exists():
        doc.add_heading("4.2. Diagramme de Cas d'Utilisation", 2)
        doc.add_paragraph(
            "Ce diagramme présente les fonctionnalités du système organisées par acteur : "
            "Client, Membre du Personnel, et Administrateur."
        )
        add_image_to_doc(
            doc,
            UML_DIR / "diagramme_cas_utilisation.png",
            "Figure 3: Diagramme de cas d'utilisation",
            width_cm=18
        )
        doc.add_paragraph()
        add_page_break(doc)
    
    # Diagrammes de séquence
    if (UML_DIR / "diagramme_sequence_reservation.png").exists():
        doc.add_heading("4.3. Diagrammes de Séquence", 2)
        
        doc.add_heading("4.3.1. Processus de Réservation", 3)
        doc.add_paragraph(
            "Le diagramme suivant illustre le processus complet de réservation d'un rendez-vous, "
            "de la sélection du service à la confirmation finale."
        )
        add_image_to_doc(
            doc,
            UML_DIR / "diagramme_sequence_reservation.png",
            "Figure 4: Diagramme de séquence - Processus de réservation",
            width_cm=18
        )
        doc.add_paragraph()
    
    if (UML_DIR / "diagramme_sequence_paiement.png").exists():
        doc.add_heading("4.3.2. Processus de Paiement", 3)
        doc.add_paragraph(
            "Ce diagramme détaille le processus de paiement avec les différentes méthodes disponibles."
        )
        add_image_to_doc(
            doc,
            UML_DIR / "diagramme_sequence_paiement.png",
            "Figure 5: Diagramme de séquence - Processus de paiement",
            width_cm=18
        )
        doc.add_paragraph()
    
    if (UML_DIR / "diagramme_sequence_reprogrammation.png").exists():
        doc.add_heading("4.3.3. Processus de Reprogrammation", 3)
        doc.add_paragraph(
            "Le diagramme suivant illustre le processus de reprogrammation d'un rendez-vous."
        )
        add_image_to_doc(
            doc,
            UML_DIR / "diagramme_sequence_reprogrammation.png",
            "Figure 6: Diagramme de séquence - Processus de reprogrammation",
            width_cm=18
        )
        doc.add_paragraph()
    
    add_page_break(doc)
    
    # ============================================
    # 5. INTERFACE UTILISATEUR
    # ============================================
    print("📝 Section 5: Interface utilisateur...")
    
    doc.add_heading("5. Interface Utilisateur", 1)
    
    doc.add_paragraph(
        "Cette section présente les différentes pages de l'application avec leurs captures d'écran."
    )
    
    # Pages publiques
    doc.add_heading("5.1. Pages Publiques", 2)
    
    pages_publiques = [
        ("01_page_accueil.png", "Page d'accueil", 
         "La page d'accueil présente la liste des services disponibles avec leurs descriptions, prix et images."),
        ("02_page_connexion.png", "Page de connexion",
         "Interface de connexion permettant aux utilisateurs de se connecter à leur compte."),
        ("03_page_inscription.png", "Page d'inscription",
         "Formulaire d'inscription pour créer un nouveau compte utilisateur."),
        ("04_page_contact.png", "Page de contact",
         "Formulaire de contact permettant aux visiteurs d'envoyer des messages."),
        ("05_nouveau_rendez_vous.png", "Page nouveau rendez-vous",
         "Page permettant de créer un nouveau rendez-vous en sélectionnant un service.")
    ]
    
    for i, (img_file, title, desc) in enumerate(pages_publiques, 1):
        img_path = SCREENSHOTS_DIR / img_file
        if img_path.exists():
            doc.add_heading(f"5.1.{i}. {title}", 3)
            doc.add_paragraph(desc)
            add_image_to_doc(
                doc,
                img_path,
                f"Figure {6 + i}: {title}",
                width_cm=16
            )
            doc.add_paragraph()
    
    add_page_break(doc)
    
    # Pages authentifiées
    doc.add_heading("5.2. Pages Authentifiées", 2)
    
    doc.add_paragraph(
        "Les pages suivantes nécessitent une authentification. Si l'utilisateur n'est pas connecté, "
        "elles redirigent automatiquement vers la page de connexion."
    )
    
    pages_auth = [
        ("06_mes_rendez_vous.png", "Mes rendez-vous",
         "Page permettant aux utilisateurs connectés de consulter et gérer leurs rendez-vous."),
        ("07_calendrier.png", "Calendrier",
         "Vue calendrier affichant tous les rendez-vous de manière visuelle."),
        ("08_modifier_profil.png", "Modifier mon profil",
         "Page permettant aux utilisateurs de modifier leurs informations personnelles."),
        ("09_changer_mot_de_passe.png", "Changer mot de passe",
         "Interface pour changer le mot de passe de son compte.")
    ]
    
    for i, (img_file, title, desc) in enumerate(pages_auth, 1):
        img_path = SCREENSHOTS_DIR / img_file
        if img_path.exists():
            doc.add_heading(f"5.2.{i}. {title}", 3)
            doc.add_paragraph(desc)
            add_image_to_doc(
                doc,
                img_path,
                f"Figure {11 + i}: {title}",
                width_cm=16
            )
            doc.add_paragraph()
    
    add_page_break(doc)
    
    # ============================================
    # 6. FONCTIONNALITÉS PRINCIPALES
    # ============================================
    print("📝 Section 6: Fonctionnalités principales...")
    
    doc.add_heading("6. Fonctionnalités Principales", 1)
    
    doc.add_heading("6.1. Gestion des Services", 2)
    doc.add_paragraph(
        "Le système permet de créer et gérer des services avec leurs caractéristiques : nom, description, "
        "durée, prix, acompte, image, devise et couleur de fond. Chaque service peut avoir des limites "
        "de reprogrammation personnalisées."
    )
    
    doc.add_heading("6.2. Gestion du Personnel", 2)
    doc.add_paragraph(
        "Les membres du personnel peuvent être configurés avec leurs horaires de travail, jours de congé, "
        "services offerts, et paramètres de disponibilité. Le système gère automatiquement les conflits "
        "et les disponibilités."
    )
    
    doc.add_heading("6.3. Système de Réservation", 2)
    doc.add_paragraph(
        "Les clients peuvent réserver des rendez-vous en ligne en sélectionnant un service, une date, "
        "une heure et un membre du personnel. Le système vérifie automatiquement les disponibilités et "
        "gère les conflits."
    )
    
    doc.add_heading("6.4. Système de Paiement", 2)
    doc.add_paragraph(
        "Le système intègre plusieurs méthodes de paiement : carte bancaire (Stripe), virement bancaire, "
        "et portefeuilles électroniques (Bankily, Masrvi, Click, Sedad, Amanty). Le paiement peut être "
        "effectué en totalité ou via un acompte."
    )
    
    doc.add_heading("6.5. Notifications et Rappels", 2)
    doc.add_paragraph(
        "Le système envoie automatiquement des emails de confirmation lors de la réservation et peut "
        "envoyer des rappels 24 heures avant le rendez-vous si configuré. Les emails incluent des "
        "fichiers ICS pour la synchronisation avec les calendriers."
    )
    
    doc.add_heading("6.6. Reprogrammation", 2)
    doc.add_paragraph(
        "Les clients peuvent reprogrammer leurs rendez-vous dans les limites définies par le service. "
        "Le système conserve un historique complet des reprogrammations."
    )
    
    add_page_break(doc)
    
    # ============================================
    # 7. CONCLUSION
    # ============================================
    print("📝 Section 7: Conclusion...")
    
    doc.add_heading("7. Conclusion", 1)
    
    conclusion_text = """
Ce système de gestion de rendez-vous représente une solution complète et moderne pour la prise de rendez-vous en ligne. 
Il offre une interface utilisateur intuitive, une gestion flexible des disponibilités, et un système de paiement intégré.

Les fonctionnalités développées permettent une gestion efficace des rendez-vous tout en offrant une expérience utilisateur 
optimale. L'architecture modulaire facilite la maintenance et l'évolution future du système.

Le système est prêt pour un déploiement en production et peut être facilement adapté aux besoins spécifiques de différents 
types d'entreprises et de professionnels.
"""
    
    doc.add_paragraph(conclusion_text.strip())
    
    # ============================================
    # ANNEXES
    # ============================================
    add_page_break(doc)
    doc.add_heading("Annexes", 1)
    
    doc.add_heading("A. Informations Techniques", 2)
    doc.add_paragraph(f"Version Django : 5.2.7")
    doc.add_paragraph(f"Version Python : 3.x")
    doc.add_paragraph(f"Base de données : SQLite/PostgreSQL")
    doc.add_paragraph(f"Date de génération du rapport : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    # Sauvegarder le document
    print()
    print("💾 Sauvegarde du document...")
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    
    print()
    print("=" * 70)
    print(f"✅ RAPPORT GÉNÉRÉ AVEC SUCCÈS !")
    print("=" * 70)
    print(f"📄 Fichier : {OUTPUT_FILE}")
    print(f"📂 Emplacement : {OUTPUT_DIR}")
    print("=" * 70)
    
    return True

if __name__ == "__main__":
    try:
        success = generate_rapport()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\n❌ Erreur : {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

